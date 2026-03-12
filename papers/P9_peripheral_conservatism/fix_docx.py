"""Post-process pandoc-generated docx to fix known conversion issues."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

doc = Document('draft_v0.1_jseas_anonymous.docx')

# ============================================================
# FIX 1: Add table numbers to table captions
# ============================================================
table_captions = {
    'PMP cognacy rates across Austronesian varieties': 'Table 1. ',
    'PMP cognacy rates (%) by semantic domain': 'Table 2. ',
    'Cross-regional mortuary structural parallels': 'Table 3. ',
}

for para in doc.paragraphs:
    if para.style and para.style.name == 'Table Caption':
        for key, prefix in table_captions.items():
            if key in para.text:
                # Prepend "Table N. " to the caption
                if not para.text.startswith('Table'):
                    for run in para.runs:
                        if key in run.text:
                            run.text = prefix + run.text
                            break
                    else:
                        # If key spans multiple runs, prepend to first run
                        if para.runs:
                            para.runs[0].text = prefix + para.runs[0].text
                print(f'  Fixed table caption: {prefix}{key[:50]}...')
                break

# ============================================================
# FIX 2: Add figure numbers to figure captions
# ============================================================
figure_captions_ordered = [
    ('Peripheral Conservatism Framework: three overwriting', 'Figure 1. '),
    ('PMP cognacy gradient across Austronesian varieties', 'Figure 2. '),
    ('PMP cognacy rates by semantic domain across four', 'Figure 3. '),
    ('Indianization wave', 'Figure 4. '),
    ('four-layer botanical palimpsest', 'Figure 5. '),
    ('Material culture in 268 Old Javanese inscriptions', 'Figure 6. '),
]

for para in doc.paragraphs:
    if para.style and para.style.name == 'Image Caption':
        for key, prefix in figure_captions_ordered:
            if key in para.text:
                if not para.text.startswith('Figure'):
                    if para.runs:
                        para.runs[0].text = prefix + para.runs[0].text
                print(f'  Fixed figure caption: {prefix}{key[:50]}...')
                break

# ============================================================
# FIX 3: Fix Table 3 (mortuary) — replace empty checkmark cells
# ============================================================
# Table 2 in 0-indexed = Table 3 in paper
if len(doc.tables) >= 3:
    mortuary_table = doc.tables[2]
    print(f'\n  Fixing mortuary table ({len(mortuary_table.rows)} rows x {len(mortuary_table.columns)} cols)')

    # Expected data (from LaTeX source):
    # Row 0: header
    # Row 1: Exposed/above-ground phase: check, check, check
    # Row 2: Secondary bone treatment: check, check, check
    # Row 3: Aromatic botanical element: check, ?, check
    # Row 4: Community participation: check, check, check
    # Row 5: Textile wrapping of bones: partial, check, check
    # Row 6: Sacred tree association: check, check, check

    checkmark_data = [
        # (row_idx, col_idx_Trunyan, col_idx_Toraja, col_idx_Malagasy)
        # Trunyan=col1, Toraja=col2, Malagasy=col3
        (1, ['✓', '✓', '✓']),           # Exposed/above-ground
        (2, ['✓', '✓', '✓']),           # Secondary bone treatment
        (3, ['✓', '?', '✓']),           # Aromatic botanical element
        (4, ['✓', '✓', '✓']),           # Community participation
        (5, ['partial', '✓', '✓']),     # Textile wrapping
        (6, ['✓', '✓', '✓']),           # Sacred tree association
    ]

    for row_idx, values in checkmark_data:
        if row_idx < len(mortuary_table.rows):
            row = mortuary_table.rows[row_idx]
            for col_offset, val in enumerate(values):
                cell_idx = col_offset + 1  # skip Element column
                if cell_idx < len(row.cells):
                    cell = row.cells[cell_idx]
                    current = cell.text.strip()
                    if current == '' or current != val:
                        # Clear existing content and set new
                        for p in cell.paragraphs:
                            p.text = ''
                        cell.paragraphs[0].text = val
                        # Center align
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if current != val:
                            print(f'    Row {row_idx}, Col {cell_idx}: "{current}" → "{val}"')

# ============================================================
# FIX 4: Ensure proper figure sizing (avoid overflow)
# ============================================================
from docx.oxml.ns import qn

img_count = 0
for para in doc.paragraphs:
    for run in para.runs:
        drawings = run._element.findall(qn('w:drawing'))
        for drawing in drawings:
            img_count += 1
            # Find the extent element to check/fix dimensions
            extents = drawing.findall('.//' + qn('wp:extent'))
            for ext in extents:
                cx = int(ext.get('cx', 0))
                cy = int(ext.get('cy', 0))
                # Convert EMU to inches (914400 EMU per inch)
                w_inches = cx / 914400
                h_inches = cy / 914400

                # Max width for A4 with 2.5cm margins = 21 - 5 = 16 cm = 6.3 inches
                max_width = 6.0  # inches, with some margin
                if w_inches > max_width:
                    scale = max_width / w_inches
                    new_cx = int(cx * scale)
                    new_cy = int(cy * scale)
                    ext.set('cx', str(new_cx))
                    ext.set('cy', str(new_cy))
                    print(f'  Image {img_count}: resized from {w_inches:.1f}" to {max_width:.1f}" wide')

                    # Also fix inline extent if present
                    inline_extents = drawing.findall('.//' + qn('a:ext'))
                    for ie in inline_extents:
                        ie_cx = int(ie.get('cx', 0))
                        ie_cy = int(ie.get('cy', 0))
                        if ie_cx == cx:  # match the outer extent
                            ie.set('cx', str(new_cx))
                            ie.set('cy', str(new_cy))
                else:
                    print(f'  Image {img_count}: OK ({w_inches:.1f}" x {h_inches:.1f}")')

# ============================================================
# FIX 5: Set document-wide font to Times New Roman 12pt
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# ============================================================
# FIX 6: Fix heading styles — remove blue, set Times New Roman
# ============================================================
from docx.shared import RGBColor

# JSEAS Style Sheet:
#   - Headings: bold, 12pt, unnumbered
#   - Sub-headings: italic, 12pt
#   - Title: bold, 14pt (centered)
#   - All Times New Roman, black

heading_config = {
    'Heading 1': {'size': Pt(12), 'bold': True, 'italic': False},   # Main section
    'Heading 2': {'size': Pt(12), 'bold': False, 'italic': True},   # Sub-section
    'Heading 3': {'size': Pt(12), 'bold': False, 'italic': True},   # Sub-sub
    'Heading 4': {'size': Pt(12), 'bold': False, 'italic': True},   # Sub-sub-sub
    'Title': {'size': Pt(14), 'bold': True, 'italic': False},       # Paper title
}

for style_name, cfg in heading_config.items():
    try:
        h_style = doc.styles[style_name]
        h_font = h_style.font
        h_font.name = 'Times New Roman'
        h_font.size = cfg['size']
        h_font.bold = cfg['bold']
        h_font.italic = cfg['italic']
        h_font.color.rgb = RGBColor(0, 0, 0)
        print(f'  Fixed style: {style_name} → TNR {cfg["size"].pt}pt, '
              f'{"bold" if cfg["bold"] else "italic"}, black')
    except KeyError:
        pass

# Fix inline heading runs (override any explicit formatting from pandoc)
for para in doc.paragraphs:
    if para.style and para.style.name.startswith('Heading'):
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = Pt(12)
            if para.style.name == 'Heading 1':
                run.font.bold = True
                run.font.italic = False
            else:  # Heading 2, 3, 4 = italic per JSEAS
                run.font.bold = False
                run.font.italic = True

# Fix title runs
for para in doc.paragraphs:
    if para.style and para.style.name == 'Title':
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.italic = False

print('  Fixed headings: H1=bold 12pt, H2/H3=italic 12pt, Title=bold 14pt (JSEAS style)')

# ============================================================
# FIX 7: Double spacing + left-aligned (JSEAS requirement)
# ============================================================
from docx.shared import Twips
from docx.enum.text import WD_LINE_SPACING

# Set double spacing on Normal style
normal_style = doc.styles['Normal']
normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Apply to all paragraphs (override any justified alignment from pandoc)
for para in doc.paragraphs:
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Double spacing for body text, keep single for captions/footnotes
    if para.style and para.style.name in ('Normal', 'Body Text', 'First Paragraph'):
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    elif para.style and para.style.name.startswith('Heading'):
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

print('  Fixed: double spacing, left-aligned (JSEAS style)')

# ============================================================
# FIX 8: Remove References/Bibliography section
# ============================================================
# JSEAS does not publish bibliographies — all refs must be in footnotes only.
# pandoc --citeproc generates Bibliography-styled paragraphs that we need to remove.
print('\n  Removing References/Bibliography section...')
paras_to_remove = []
for i, para in enumerate(doc.paragraphs):
    if para.style and para.style.name == 'Bibliography':
        paras_to_remove.append(para)

if paras_to_remove:
    print(f'    Found {len(paras_to_remove)} Bibliography paragraphs')
    for para in paras_to_remove:
        p_element = para._element
        p_element.getparent().remove(p_element)
    print(f'    Removed {len(paras_to_remove)} paragraphs')
else:
    print('    No Bibliography paragraphs found')

# ============================================================
# SAVE
# ============================================================
output = 'draft_v0.1_jseas_anonymous.docx'
doc.save(output)
print(f'\nSaved fixed document: {output}')
print('Done!')
