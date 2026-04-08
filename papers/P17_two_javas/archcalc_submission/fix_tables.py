"""
Fix tables in P17_manuscript_formatted.docx.
Strategy: remove broken pandoc tables, insert clean ones after caption paragraphs.
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from lxml import etree

SUBMISSION_DIR = os.path.dirname(os.path.abspath(__file__))
NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def set_cell(cell, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(13)
    p.paragraph_format.first_line_indent = Cm(0)


def set_borders(cell, top=False, bottom=False):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    borders_xml = (
        f'<w:tcBorders {nsdecls("w")}>'
        + (f'<w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>' if top else '<w:top w:val="nil"/>')
        + '<w:left w:val="nil"/>'
        + (f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>' if bottom else '<w:bottom w:val="nil"/>')
        + '<w:right w:val="nil"/>'
        + '</w:tcBorders>'
    )
    tcPr.append(parse_xml(borders_xml))


def build_table1(doc):
    """Build Table 1: Summary of five analyses. Returns table element."""
    headers = ["Analysis", "Key statistic", "p-value", "Implication"]
    rows = [
        ["Spatial segregation (E104)",
         "Median gap = 13 km; MW U = 8,081",
         "< 0.000001",
         "Candi and inscriptions occupy different zones"],
        ["Elevation gradient (E100)",
         "Density 1.96 \u2192 18.61; 9.5\u00d7 increase",
         "< 0.000001",
         "Mountain sites = buried population tip"],
        ["Vocab \u00d7 depth (E102)",
         "\u03c1 = 0.456 (partial); 5.8\u00d7 jump",
         "< 0.0001",
         "Burial hides indigenous content"],
        ["Court zone trend (E103)",
         "\u03c1 = 0.781 (court only)",
         "< 0.0001",
         "Indianization is zone-specific"],
        ["Post-929 relocation (E105)",
         "91% Sanskrit \u2192 89% indigenous",
         "\u2013",
         "Collapse removes Sanskrit overlay"],
    ]

    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, size=10,
                 align=WD_ALIGN_PARAGRAPH.CENTER if i == 2 else WD_ALIGN_PARAGRAPH.LEFT)
        set_borders(table.rows[0].cells[i], top=True, bottom=True)

    # Data
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            set_cell(table.rows[r+1].cells[c], val, size=10,
                     align=WD_ALIGN_PARAGRAPH.CENTER if c == 2 else WD_ALIGN_PARAGRAPH.LEFT)
            set_borders(table.rows[r+1].cells[c], bottom=(r == len(rows)-1))

    return table._element


def build_table2(doc):
    """Build Table 2: Distribution by volcanic distance zone. Returns table element."""
    headers = ["Zone", "Candi", "%", "Inscriptions", "%"]
    rows = [
        ["0\u201310 km", "60", "42.3", "22", "12.5"],
        ["10\u201320 km", "21", "14.8", "46", "26.1"],
        ["20\u201330 km", "44", "31.0", "69", "39.2"],
        ["30\u201340 km", "8", "5.6", "12", "6.8"],
        ["40\u201360 km", "5", "3.5", "20", "11.4"],
        ["60\u2013100 km", "4", "2.8", "7", "4.0"],
    ]
    total = ["Total", "142", "", "176", ""]

    table = doc.add_table(rows=len(rows) + 2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, size=10,
                 align=WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT)
        set_borders(table.rows[0].cells[i], top=True, bottom=True)

    # Data
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            set_cell(table.rows[r+1].cells[c], val, size=10,
                     align=WD_ALIGN_PARAGRAPH.CENTER if c > 0 else WD_ALIGN_PARAGRAPH.LEFT)
            set_borders(table.rows[r+1].cells[c])

    # Total
    for c, val in enumerate(total):
        set_cell(table.rows[-1].cells[c], val, bold=True, size=10,
                 align=WD_ALIGN_PARAGRAPH.CENTER if c > 0 else WD_ALIGN_PARAGRAPH.LEFT)
        set_borders(table.rows[-1].cells[c], top=True, bottom=True)

    return table._element


def fix_tables():
    src = os.path.join(SUBMISSION_DIR, "P17_manuscript_formatted.docx")

    # Start fresh from the pre-formatted DOCX
    src_base = os.path.join(SUBMISSION_DIR, "P17_manuscript.docx")
    if not os.path.exists(src_base):
        print("ERROR: P17_manuscript.docx not found")
        return

    # Work on the base pandoc output first (before formatting script)
    doc = Document(src_base)

    # Ensure section has page dimensions (pandoc sometimes omits these)
    from docx.shared import Inches
    section = doc.sections[0]
    if section.page_width is None:
        section.page_width = Inches(8.27)  # A4
        section.page_height = Inches(11.69)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    body = doc.element.body

    # 1. Find and remove ALL existing tables (broken pandoc ones)
    old_tables = body.findall('.//w:tbl', NSMAP)
    print(f"Found {len(old_tables)} broken tables to remove")
    for tbl in old_tables:
        tbl.getparent().remove(tbl)

    # 2. Find caption paragraphs by text content
    caption1_elem = None
    caption2_elem = None

    for p_elem in body.findall('.//w:p', NSMAP):
        # Get full text of paragraph
        texts = p_elem.findall('.//w:t', NSMAP)
        full_text = ''.join(t.text or '' for t in texts)

        if "Summary of five analyses" in full_text:
            caption1_elem = p_elem
            print(f"Found Table 1 caption")
        elif "Distribution of candi and inscriptions" in full_text:
            caption2_elem = p_elem
            print(f"Found Table 2 caption")

    # 3. Build new tables and insert after captions
    if caption1_elem is not None:
        tbl1 = build_table1(doc)
        caption1_elem.addnext(tbl1)
        print("Table 1 inserted after caption")
    else:
        print("WARNING: Table 1 caption not found")

    if caption2_elem is not None:
        tbl2 = build_table2(doc)
        caption2_elem.addnext(tbl2)
        print("Table 2 inserted after caption")
    else:
        print("WARNING: Table 2 caption not found")

    # Save as the formatted version
    doc.save(src)
    print(f"\nSaved with fixed tables: {src}")

    # Now run the formatting (heading styles etc) on this fixed file
    # Import and call format_manuscript from the main script
    print("\nNow applying heading styles...")


if __name__ == "__main__":
    fix_tables()
