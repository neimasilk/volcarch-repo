"""
Format P17 submission files for Archeologia e Calcolatori.
Based on actual published ArchCalc papers (verified 2026-04-02).

Key rules (from real papers, not just guidelines):
- NO paragraph numbering (sections numbered hierarchically: 1., 1.1, 1.2)
- Title: ALL CAPITALS
- Section headings: Small caps
- Subsection headings: Italics
- Dashes: en-dash with spaces ( -- ) for punctuation (already fixed in LaTeX)
- Captions: "Fig. N -- description" format
- Bibliography: separate DOCX, Harvard, journal titles in guillemets
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SUBMISSION_DIR = os.path.dirname(os.path.abspath(__file__))


def set_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def make_small_caps(run):
    rPr = run._element.get_or_add_rPr()
    smallCaps = rPr.makeelement(qn('w:smallCaps'), {qn('w:val'): 'true'})
    rPr.append(smallCaps)


def format_manuscript():
    """Format P17_manuscript.docx with heading styles (NO paragraph numbering)."""
    src = os.path.join(SUBMISSION_DIR, "P17_manuscript.docx")
    dst = os.path.join(SUBMISSION_DIR, "P17_manuscript_formatted.docx")

    doc = Document(src)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""

        # Title -> ALL CAPITALS
        if style_name.startswith("Title"):
            for run in para.runs:
                run.text = run.text.upper()
                set_font(run, bold=True, size=14)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # Heading 1 -> Small caps (section headings like "1. Introduction")
        if style_name == "Heading 1":
            for run in para.runs:
                set_font(run, bold=False, size=12)
                make_small_caps(run)
            continue

        # Heading 2 -> Italics (subsection headings like "2.1 Volcanic geography")
        if style_name == "Heading 2":
            for run in para.runs:
                set_font(run, bold=False, italic=True, size=12)
            continue

        # Heading 3+ -> Regular
        if style_name.startswith("Heading"):
            for run in para.runs:
                set_font(run, bold=False, italic=False, size=12)
            continue

        # Fix caption format: "Figure N:" -> "Fig. N --"
        if text.startswith("Figure ") and ":" in text[:20]:
            for run in para.runs:
                run.text = re.sub(
                    r'^Figure (\d+):',
                    r'Fig. \1 --',
                    run.text
                )

    # Set document-wide formatting
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = Pt(24)
    pf.first_line_indent = Cm(1)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(dst)
    print(f"Formatted manuscript saved: {dst}")


def create_bibliography_docx():
    """Convert P17_bibliography.txt to properly formatted .docx."""
    src = os.path.join(SUBMISSION_DIR, "P17_bibliography.txt")
    dst = os.path.join(SUBMISSION_DIR, "P17_bibliography.docx")

    with open(src, "r", encoding="utf-8") as f:
        lines = f.read().strip().split("\n")

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = Pt(24)

    title = doc.add_paragraph()
    run = title.add_run("BIBLIOGRAPHY")
    set_font(run, bold=True, size=14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    current_entry = []
    for line in lines[2:]:
        if line.strip() == "":
            if current_entry:
                entry_text = " ".join(current_entry)
                p = doc.add_paragraph()
                run = p.add_run(entry_text)
                set_font(run, size=12)
                p.paragraph_format.first_line_indent = Cm(-1)
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_after = Pt(6)
                current_entry = []
        else:
            current_entry.append(line.strip())

    if current_entry:
        entry_text = " ".join(current_entry)
        p = doc.add_paragraph()
        run = p.add_run(entry_text)
        set_font(run, size=12)
        p.paragraph_format.first_line_indent = Cm(-1)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(6)

    doc.save(dst)
    print(f"Bibliography DOCX saved: {dst}")


def create_captions_docx():
    """Convert P17_figure_captions to .docx with ArchCalc format (Fig. N --)."""
    src = os.path.join(SUBMISSION_DIR, "P17_figure_captions.docx.txt")
    dst = os.path.join(SUBMISSION_DIR, "P17_figure_captions.docx")

    with open(src, "r", encoding="utf-8") as f:
        content = f.read().strip()

    # Fix caption format: "Figure N." -> "Fig. N --"
    content = re.sub(r'Figure (\d+)\.', r'Fig. \1 --', content)
    content = re.sub(r'Table (\d+)\.', r'Tab. \1 --', content)

    lines = content.split("\n")

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = Pt(24)

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line == "FIGURE CAPTIONS":
            p = doc.add_paragraph()
            run = p.add_run("FIGURE CAPTIONS")
            set_font(run, bold=True, size=14)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            continue

        if line == "TABLE CAPTIONS":
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("TABLE CAPTIONS")
            set_font(run, bold=True, size=14)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            continue

        p = doc.add_paragraph()
        run = p.add_run(line)
        set_font(run, size=12)
        p.paragraph_format.space_after = Pt(12)

    doc.save(dst)
    print(f"Figure captions DOCX saved: {dst}")


def run_compliance_audit():
    """Run full compliance check against ArchCalc requirements."""
    tex_file = os.path.join(SUBMISSION_DIR, "..", "draft_v0.3_archcalc.tex")
    with open(tex_file, "r", encoding="utf-8") as f:
        text = f.read()

    print("\n=== ArchCalc Compliance Audit ===\n")

    # 1. Em-dashes (should be 0)
    em_count = text.count('---')
    status = "PASS" if em_count == 0 else "FAIL"
    print(f"[{status}] Em-dashes (---): {em_count} (should be 0)")

    # 2. Figure references
    figure_full = len(re.findall(r'\\b[Ff]igure\\b', text)) + text.count('Figure')
    fig_abbr = text.count('Fig.')
    status = "PASS" if figure_full == 0 else "WARN"
    print(f"[{status}] 'Figure' in text: {figure_full}, 'Fig.': {fig_abbr}")

    # 3. Footnotes
    footnotes = text.count('\\footnote{')
    status = "PASS" if footnotes == 0 else "FAIL"
    print(f"[{status}] Footnotes: {footnotes} (should be 0)")

    # 4. Anonymization
    author_refs = text.lower().count('amien') + text.lower().count('volcarch')
    status = "PASS" if author_refs == 0 else "FAIL"
    print(f"[{status}] Author name/VOLCARCH refs: {author_refs} (should be 0 for double-blind)")

    # 5. Abstract word count
    abs_match = re.search(r'\\begin\{abstract\}(.*?)\\end\{abstract\}', text, re.DOTALL)
    if abs_match:
        abs_text = re.sub(r'\\[a-zA-Z]+\{[^}]*\}|\\[a-zA-Z]+|\$[^$]*\$|~|``|\'\'', ' ', abs_match.group(1))
        abs_words = len(abs_text.split())
        status = "PASS" if abs_words <= 200 else "FAIL"
        print(f"[{status}] Abstract words: ~{abs_words} (max 200)")

    # 6. Spelling consistency
    british = {'artefact': 0, 'centre': 0, 'civilisation': 0, 'characterised': 0, 'organisation': 0}
    american = {'artifact': 0, 'center': 0, 'civilization': 0, 'characterized': 0, 'organization': 0}
    text_lower = text.lower()

    br_found = {k: text_lower.count(k) for k in british if text_lower.count(k) > 0}
    am_found = {k: text_lower.count(k) for k in american if text_lower.count(k) > 0}

    # Filter out LaTeX commands for American spellings
    real_american = {}
    for word, count in am_found.items():
        # Check if the word appears outside LaTeX commands
        pattern = r'(?<!\\)(?<!\{)\b' + word + r'\b'
        real_matches = len(re.findall(pattern, text_lower))
        if real_matches > 0:
            real_american[word] = real_matches

    if br_found and not real_american:
        print(f"[PASS] Spelling: British English consistent ({br_found})")
    elif real_american:
        print(f"[FAIL] Mixed spelling: British={br_found}, American={real_american}")
    else:
        print(f"[PASS] No British/American variants detected")

    # 7. Section numbering (LaTeX \section auto-numbers)
    sections = len(re.findall(r'\\section\{', text))
    subsections = len(re.findall(r'\\subsection\{', text))
    print(f"[INFO] Sections: {sections}, Subsections: {subsections} (LaTeX auto-numbers these)")

    # 8. Figure count
    figures = len(re.findall(r'\\begin\{figure\}', text))
    tables = len(re.findall(r'\\begin\{table\}', text))
    total = figures + tables
    status = "PASS" if total <= 10 else "FAIL"
    print(f"[{status}] Figures: {figures}, Tables: {tables}, Total: {total} (max 10)")

    # 9. Check submission files exist
    print(f"\n--- Submission Files ---")
    required = [
        "P17_manuscript_formatted.docx",
        "P17_bibliography.docx",
        "P17_figures.zip",
        "P17_figure_captions.docx"
    ]
    for f in required:
        path = os.path.join(SUBMISSION_DIR, f)
        exists = os.path.exists(path)
        size = os.path.getsize(path) if exists else 0
        status = "PASS" if exists and size > 0 else "FAIL"
        print(f"[{status}] {f} ({size:,} bytes)")


if __name__ == "__main__":
    print("=" * 60)
    print("P17 ArchCalc Submission Formatter (v2)")
    print("Based on actual published ArchCalc papers")
    print("=" * 60)

    print("\n--- Step 1: Format manuscript (NO paragraph numbering) ---")
    format_manuscript()

    print("\n--- Step 2: Create bibliography DOCX ---")
    create_bibliography_docx()

    print("\n--- Step 3: Create figure captions DOCX (Fig. N -- format) ---")
    create_captions_docx()

    print("\n--- Step 4: Compliance audit ---")
    run_compliance_audit()

    print("\n" + "=" * 60)
    print("DONE.")
    print("=" * 60)
